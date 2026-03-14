# indexer.py v3 — Small-to-Big チャンク戦略・メタデータ強制スキーマ対応
#
# 変更履歴:
#   v3 (2026-02-25): Small-to-Big チャンク、doc_type自動判定、
#                    EXCLUDE_PATTERNS対応、メタデータバリデーション追加

import os
import json
import logging
import hashlib
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple

import fitz  # PyMuPDF
from docx import Document
from google.genai import types
import chromadb
from chromadb import EmbeddingFunction

from classifier import DocumentClassifier
from chunk_builder import ChunkBuilder
from dense_indexer import DenseIndexer
from lexical_indexer import LexicalIndexer
from metadata_repository import MetadataRepository

logger = logging.getLogger(__name__)

from config import (
    KNOWLEDGE_BASE_DIR,
    SEARCH_MD_DIR,   # = KNOWLEDGE_BASE_DIR (後方互換エイリアス)
    REFERENCE_DIR,   # = KNOWLEDGE_BASE_DIR (後方互換エイリアス)
    CHROMA_DB_DIR,
    FILE_INDEX_PATH,
    SUPPORTED_EXTENSIONS,
    COLLECTION_NAME,
    EXCLUDE_FOLDERS,
    EMBEDDING_MODEL,
    PARENT_CHUNKS_DIR,
)
from dense_indexer import get_chroma_client
from gemini_client import get_client
from utils.retry import sync_retry

# ─── チャンクサイズ設定（doc_type 別） ──────────────────────────────────────────
# 小チャンク（ChromaDB に登録する検索用チャンク）
SMALL_CHUNK_SIZES: Dict[str, int] = {
    "catalog":  150,
    "drawing":  100,
    "spec":     200,
    "law":      200,
    "_default": 200,
}
CHUNK_OVERLAP: int = 30

# 親チャンク（LLM へ渡すコンテキスト用）
PARENT_CHUNK_MIN: int = 500
PARENT_CHUNK_MAX: int = 800

# (get_chroma_client was moved to dense_indexer.py)


# ─── Embedding ─────────────────────────────────────────────────────────────────
@sync_retry(max_retries=3, base_wait=2.0)
def _call_embed_content(client, model, contents, config):
    return client.models.embed_content(
        model=model,
        contents=contents,
        config=config
    )


class GeminiEmbeddingFunction(EmbeddingFunction):
    """Gemini Embedding API を使用するカスタム EmbeddingFunction"""

    def __call__(self, input: List[str]) -> List[List[float]]:
        client = get_client()
        # Batch size for Gemini is 100
        batch_size = 100
        embeddings = []
        for i in range(0, len(input), batch_size):
            batch = input[i : i + batch_size]
            try:
                result = _call_embed_content(
                    client,
                    model=EMBEDDING_MODEL,
                    contents=batch,
                    config=types.EmbedContentConfig(task_type="retrieval_document")
                )
                for emb in result.embeddings:
                    embeddings.append(emb.values)
            except Exception as e:
                logger.error(f"Batch embedding error: {e}")
                raise e
        return embeddings


def get_query_embedding(text: str) -> List[float]:
    """検索クエリ用の Embedding 取得"""
    client = get_client()
    result = _call_embed_content(
        client,
        model=EMBEDDING_MODEL,
        contents=text,
        config=types.EmbedContentConfig(task_type="retrieval_query")
    )
    return result.embeddings[0].values


# (GeminiEmbeddingFunction and get_query_embedding remain for backward compatibility if needed, 
# but DenseIndexer is preferred for new code)

# ─── doc_type 自動判定 (classifier.py に移行済み。互換性のためのラッパー) ──────────
def _infer_doc_type(category: str, filename: str) -> str:
    """カテゴリとファイル名から doc_type を推定 (DocumentClassifierに委譲)"""
    return DocumentClassifier().infer_doc_type(category, filename)


# ─── ファイルスキャン ────────────────────────────────────────────────────────────
def _should_exclude(path: Path, base_path: Path) -> bool:
    """除外ルールに該当するか判定"""
    try:
        from config import EXCLUDE_PATTERNS
    except ImportError:
        EXCLUDE_PATTERNS = []
    rel = str(path.relative_to(base_path))
    # フォルダ除外
    if any(ex in rel for ex in EXCLUDE_FOLDERS):
        return True
    # ファイル名パターン除外
    if any(p in path.name for p in EXCLUDE_PATTERNS):
        return True
    return False


def scan_files(base_dir: Path = SEARCH_MD_DIR) -> List[Dict[str, Any]]:
    """ナレッジ DB フォルダを再帰スキャンしてファイルメタデータを収集"""
    files = []
    base_path = Path(base_dir)
    if not base_path.exists():
        base_path.mkdir(parents=True, exist_ok=True)
        return files

    raw_files = []
    for filepath in base_path.rglob("*"):
        if filepath.is_dir():
            continue
        if _should_exclude(filepath, base_path):
            continue
        if filepath.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue

        rel_path = filepath.relative_to(base_path)
        parts = rel_path.parts
        category = parts[0] if len(parts) > 1 else ""
        subcategory = parts[1] if len(parts) > 2 else ""
        sub_subcategory = parts[2] if len(parts) > 3 else ""
        doc_type = _infer_doc_type(category, filepath.name)

        raw_files.append({
            "filename":       filepath.name,
            "full_path":      str(filepath),
            "rel_path":       str(rel_path),
            "category":       category,
            "subcategory":    subcategory,
            "sub_subcategory": sub_subcategory,
            "file_type":      filepath.suffix.lower().lstrip('.'),
            "file_size_kb":   round(filepath.stat().st_size / 1024, 2),
            "modified_at":    datetime.fromtimestamp(filepath.stat().st_mtime, tz=timezone.utc).isoformat(),
            "doc_type":       doc_type,
        })

    # MD 優先ロジック: 同フォルダに同名 PDF がある場合は PDF をスキップ
    processed_paths = set()
    final_files = []
    for f in raw_files:
        if f["file_type"] == "md":
            pdf_rel = str(Path(f["rel_path"]).with_suffix(".pdf"))
            pdf_full = base_path / pdf_rel
            if pdf_full.exists():
                f["source_pdf_rel"] = pdf_rel
                processed_paths.add(pdf_rel)
            final_files.append(f)
            processed_paths.add(f["rel_path"])
    for f in raw_files:
        if f["rel_path"] not in processed_paths:
            final_files.append(f)

    return final_files


def scan_pdfs_dir() -> List[Dict[str, Any]]:
    """data/pdfs/ 内の全 PDF をスキャン（再インデックス用）"""
    pdfs = []
    pdf_dir = Path(PARENT_CHUNKS_DIR).parent / "pdfs"  # BASE_DIR/pdfs
    if not pdf_dir.exists():
        return []
    for f in pdf_dir.glob("*.pdf"):
        if f.stat().st_size < 10:
            logger.warning(f"Skipping too-small PDF (possibly corrupt): {f.name}")
            continue
        pdfs.append({
            "full_path":       str(f),
            "source_pdf_hash": f.stem,  # ハッシュ名がファイル名
            "filename":        f.name,
        })
    return pdfs


# ─── テキスト抽出 ────────────────────────────────────────────────────────────────
def extract_text(filepath: str) -> List[Dict[str, Any]]:
    """ファイルからテキストを抽出（戻り値: [{text, page_number, has_image}, ...]）"""
    path = Path(filepath)
    ext = path.suffix.lower()
    try:
        if ext == ".pdf":
            return _extract_pdf(filepath)
        elif ext in [".md", ".txt"]:
            return _extract_text_file(filepath)
        elif ext == ".docx":
            return _extract_docx(filepath)
        return []
    except (fitz.FileDataError, ValueError, IOError) as e:
        logger.error(f"  PyMuPDF open error {filepath}: {e}")
        return []


def _extract_pdf(filepath: str) -> List[Dict[str, Any]]:
    """PDF からテキストとページ情報を抽出"""
    pages = []
    doc = fitz.open(filepath)
    for page_num, page in enumerate(doc, start=1):
        text = page.get_text()
        # 画像の有無を判定（ページ内の画像オブジェクト数）
        image_list = page.get_images(full=False)
        has_image = len(image_list) > 0
        if text.strip():
            pages.append({
                "text":       text,
                "page_number": page_num,
                "has_image":  has_image,
            })
    doc.close()
    return pages


def _extract_text_file(filepath: str) -> List[Dict[str, Any]]:
    """Markdown / テキストファイルを抽出（Frontmatter を除去）"""
    with open(filepath, "r", encoding="utf-8") as f:
        text = f.read()
    if text.startswith("---"):
        parts = text.split("---", 2)
        if len(parts) >= 3:
            text = parts[2].lstrip("\n")
    return [{"text": text, "page_number": None, "has_image": False}] if text.strip() else []


def _extract_docx(filepath: str) -> List[Dict[str, Any]]:
    doc = Document(filepath)
    text = "\n".join(p.text for p in doc.paragraphs)
    return [{"text": text, "page_number": None, "has_image": False}] if text.strip() else []


def parse_frontmatter(filepath: str) -> Dict[str, Any]:
    """YAML Frontmatter を解析して辞書を返す"""
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            content = f.read()
        if content.startswith("---"):
            import yaml
            parts = content.split("---", 2)
            if len(parts) >= 3:
                fm = yaml.safe_load(parts[1])
                return fm if isinstance(fm, dict) else {}
    except (Exception) as e:
        logger.warning(f"Frontmatter parsing error ({filepath}): {e}")
    return {}


# ─── Small-to-Big チャンク ───────────────────────────────────────────────────────
def _split_into_small_chunks(text: str, size: int, overlap: int) -> List[str]:
    """固定サイズの小チャンクに分割（オーバーラップ付き）"""
    if not text.strip():
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + size, len(text))
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = end - overlap
    return chunks


def _split_into_parent_chunks(text: str, min_size: int = PARENT_CHUNK_MIN, max_size: int = PARENT_CHUNK_MAX) -> List[str]:
    """段落境界を尊重しながら親チャンクに分割"""
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    parents = []
    current = ""
    for para in paragraphs:
        candidate = (current + "\n\n" + para).strip() if current else para
        if len(candidate) <= max_size:
            current = candidate
        else:
            if current:
                parents.append(current)
            # パラグラフ自体が max_size を超える場合は強制分割
            if len(para) > max_size:
                for i in range(0, len(para), max_size):
                    parents.append(para[i:i + max_size])
                current = ""
            else:
                current = para
    if current:
        parents.append(current)
    return parents


def save_parent_chunk(source_pdf_hash: str, parent_id: str, text: str) -> str:
    """親チャンク MD ファイルを保存し、ファイルパスを返す"""
    parent_dir = Path(PARENT_CHUNKS_DIR) / source_pdf_hash
    parent_dir.mkdir(parents=True, exist_ok=True)
    parent_path = parent_dir / f"{parent_id}.md"
    parent_path.write_text(text, encoding="utf-8")
    return str(parent_path)


def load_parent_chunk(parent_chunk_id: str) -> Optional[str]:
    """parent_chunk_id から親チャンクテキストを取得
    parent_chunk_id 形式: "{source_pdf_hash}/{parent_id}"
    """
    try:
        parts = parent_chunk_id.split("/", 1)
        if len(parts) != 2:
            return None
        pdf_hash, pid = parts
        parent_path = Path(PARENT_CHUNKS_DIR) / pdf_hash / f"{pid}.md"
        if parent_path.exists():
            return parent_path.read_text(encoding="utf-8")
    except Exception as e:
        logger.warning(f"Failed to load parent_chunk ({parent_chunk_id}): {e}")
    return None


def chunk_for_indexing(
    text: str,
    page_number: Optional[int],
    has_image: bool,
    doc_type: str,
    source_pdf_hash: str,
    source_pdf_name: str,
    category: str,
    rel_path: str,
    filename: str,
    file_type: str,
    modified_at: str,
    tags_str: str = "",
    drawing_type: Optional[str] = None,
    scale: Optional[str] = None,
    drive_file_id: Optional[str] = None,
    version_id: str = "legacy"
) -> List[Dict[str, Any]]:
    """
    [Legacy Wrapper] 
    新しい ChunkBuilder は Markdown 全体を受け取るため、この関数は内部で
    簡略化されたメタデータ付与を行う。
    """
    # 互換性のための簡易実装
    from chunk_builder import ChunkBuilder
    builder = ChunkBuilder()
    
    # 単一ページ/ブロックを処理する場合のダミー OCR 形式
    ocr_results = [{"text": text, "start_page": page_number or 1, "label": f"Page {page_number or 1}"}]
    source_metadata = {
        "source_pdf_hash": source_pdf_hash,
        "source_pdf_name": source_pdf_name,
        "category": category,
        "doc_type": doc_type,
        "rel_path": rel_path,
        "filename": filename,
        "file_type": file_type,
        "version_id": version_id,
        "tags_str": tags_str,
        "drive_file_id": drive_file_id or ""
    }
    
    chunks = builder.build(text, ocr_results, source_metadata)
    
    # 旧形式 (ChromaDB 内部メタデータ) に変換
    output = []
    for c in chunks:
        output.append({
            "text": c["content"],
            "metadata": {**c["metadata"], "chunk_id": c["id"]}
        })
    return output


# ─── ID 生成 ─────────────────────────────────────────────────────────────────────
def generate_doc_id(source_pdf_hash: str, rel_path: str, chunk_index: int) -> str:
    key = f"{source_pdf_hash}::{rel_path}::{chunk_index}"
    return hashlib.md5(key.encode()).hexdigest()


# ─── DB 操作 ─────────────────────────────────────────────────────────────────────
def _upsert_doc_index(rel_path: str, file_info: Dict[str, Any], chunk_count: int):
    from metadata_repository import MetadataRepository
    from database import get_session, LegacyDocument
    repo = MetadataRepository()
    
    # 1. Update Legacy Status for historical compatibility
    session = get_session()
    try:
        doc = session.query(LegacyDocument).filter(LegacyDocument.file_path == rel_path).first()
        if not doc:
            doc = LegacyDocument(
                filename=Path(rel_path).name,
                file_path=rel_path,
                file_type=Path(rel_path).suffix.lower().lstrip("."),
            )
            session.add(doc)
        
        doc.file_hash    = file_info.get("modified_at", "")
        doc.chunk_count  = chunk_count
        doc.last_indexed_at = datetime.now(timezone.utc)
        doc.file_size    = int(file_info.get("file_size_kb", 0) * 1024)
        doc.category     = file_info.get("category", "")
        doc.subcategory  = file_info.get("subcategory", "")
        doc.doc_type     = file_info.get("doc_type", "")
        doc.source_pdf_hash = file_info.get("source_pdf_hash", "")
        doc.source_pdf_name = file_info.get("source_pdf_name", "")
        doc.drive_file_id = file_info.get("drive_file_id", "")
        doc.status       = "completed"
        doc.updated_at   = datetime.now(timezone.utc)
        session.commit()
    except Exception as e:
        session.rollback()
        logger.error(f"Legacy DB upsert error for {rel_path}: {e}")
    finally:
        session.close()

    # 2. Update Unified Metadata via Repository
    try:
        # Create or update document version entry
        repo.create_document_version(
            filename=file_info.get("filename", ""),
            file_path=rel_path,
            source_pdf_hash=file_info.get("source_pdf_hash", ""),
            file_size=int(file_info.get("file_size_kb", 0) * 1024)
        )
        # Mark as searchable
        repo.mark_as_searchable(rel_path)
    except Exception as e:
        logger.error(f"Unified Metadata update error for {rel_path}: {e}")


def _delete_doc_index(rel_path: str) -> bool:
    """DBからインデックス情報を削除"""
    from database import get_session, LegacyDocument, Document as DbDocument
    session = get_session()
    try:
        # Legacy
        doc = session.query(LegacyDocument).filter(LegacyDocument.file_path == rel_path).first()
        if doc:
            session.delete(doc)
        
        # New
        new_doc = session.query(DbDocument).filter(DbDocument.file_path == rel_path).first()
        if new_doc:
            session.delete(new_doc)
            
        session.commit()
        return True
    except Exception as e:
        session.rollback()
        logger.error(f"DB deletion error for {rel_path}: {e}")
        return False
    finally:
        session.close()


def process_and_index_file(
    file_info: Dict[str, Any],
    stats: Dict[str, int],
) -> bool:
    """単一ファイル情報を受け取ってインデックスする（統一パイプライン）"""
    rel_path  = file_info["rel_path"]
    full_path = file_info["full_path"]
    
    classifier = DocumentClassifier()
    doc_type = classifier.infer_doc_type(file_info.get("category", ""), file_info.get("filename", ""))
    file_info["doc_type"] = doc_type

    pages = extract_text(full_path)
    if not pages:
        stats["skipped"] += 1
        return False

    frontmatter = parse_frontmatter(full_path) if file_info["file_type"] == "md" else {}

    # source_pdf_hash 解決
    source_pdf_hash = (
        frontmatter.get("source_pdf")
        or file_info.get("source_pdf_rel", "")
        or file_info.get("source_pdf_hash", "")
    )
    if not source_pdf_hash:
        try:
            source_pdf_hash = hashlib.sha256(Path(full_path).read_bytes()).hexdigest()[:16]
        except Exception:
            source_pdf_hash = "unknown"
    
    file_info["source_pdf_hash"] = source_pdf_hash
    source_pdf_name = (
        frontmatter.get("pdf_filename") or frontmatter.get("source_pdf_name") or file_info.get("filename", "")
    )
    file_info["source_pdf_name"] = source_pdf_name

    # MetadataRepository を通じて初期登録 (Status: accepted -> processing)
    repo = MetadataRepository()
    ids = repo.create_document_version(
        filename=file_info["filename"],
        file_path=rel_path,
        source_pdf_hash=source_pdf_hash,
        file_size=int(file_info.get("file_size_kb", 0) * 1024)
    )
    version_id = ids["version_id"]
    repo.update_ingest_stage_by_version_id(version_id, "indexing", total_pages=len(pages))

    # チャンク生成 (Unified ChunkBuilder)
    builder = ChunkBuilder()
    ocr_results = []
    full_text = ""
    for p in pages:
        ocr_results.append({
            "text": p["text"],
            "start_page": p.get("page_number", 0),
            "label": f"Page {p.get('page_number', 0)}"
        })
        full_text += p["text"] + "\n\n"

    source_metadata = {
        "version_id": version_id,
        "source_pdf_hash": source_pdf_hash,
        "source_pdf_name": source_pdf_name,
        "rel_path": rel_path,
        "filename": file_info["filename"],
        "category": frontmatter.get("primary_category") or file_info["category"],
        "doc_type": doc_type,
        "tags_str": ",".join(frontmatter.get("tags", []))
    }

    chunks = builder.build(full_text, ocr_results, source_metadata)
    
    # インデックス登録
    dense = DenseIndexer()
    dense.upsert_chunks(version_id, chunks)
    
    lexical = LexicalIndexer()
    lexical.upsert_chunks(version_id, chunks)

    # 完了更新
    _upsert_doc_index(rel_path, file_info, len(chunks)) # Legacy DB Sync
    repo.mark_as_searchable(rel_path)
    
    stats["indexed"] += 1
    stats["chunks"] += len(chunks)
    return True


# ─── メインインデックス構築 ─────────────────────────────────────────────────────
def build_index(force_rebuild: bool = False) -> Dict[str, int]:
    """knowledge_base をスキャンして ChromaDB/SQLite にインデックスを構築"""
    from database import init_db
    init_db()

    logger.info("ファイルをスキャン中...")
    files = scan_files()
    logger.info(f"発見したファイル数: {len(files)}")

    indexed_files = load_file_index().get("files", {})
    stats = {"total_files": len(files), "indexed": 0, "skipped": 0, "errors": 0, "chunks": 0}

    from metadata_repository import MetadataRepository
    repo = MetadataRepository()

    for file_info in files:
        rel_path = file_info["rel_path"]
        
        # 1. 処理中チェック
        db_status = repo.get_document_status(rel_path)
        if db_status and db_status.get("status") in ("processing", "ocr_processing", "indexing", "classified"):
            logger.info(f"Skipping {rel_path}: already being processed by ingestion pipeline.")
            stats["skipped"] += 1
            continue

        # 2. 変更検出 (force_rebuild でない場合)
        if not force_rebuild and rel_path in indexed_files:
            if indexed_files[rel_path].get("modified_at") == file_info["modified_at"]:
                stats["skipped"] += 1
                continue

        try:
            process_and_index_file(file_info, stats)
        except Exception as e:
            logger.error(f"インデックスエラー ({rel_path}): {e}", exc_info=True)
            stats["errors"] += 1

    logger.info(f"インデックス完了: {stats['indexed']}ファイル, {stats['chunks']}チャンク")
    return stats


# ─── 全件再インデックス (data/pdfs/ から) ──────────────────────────────────────
def reindex_from_pdfs(progress_interval: int = 10) -> Dict[str, Any]:
    """
    data/pdfs/ 内の全 PDF を対象に取り込みパイプラインを実行。
    各 PDF は OCR 済み MD が parent_chunks/ にあれば使用し、
    なければ直接テキスト抽出してインデックスする。
    """
    from database import init_db
    init_db()

    pdfs = scan_pdfs_dir()
    total = len(pdfs)
    logger.info(f"data/pdfs/ から {total} 件の PDF を再インデックス開始")

    stats = {"total": total, "indexed": 0, "skipped": 0, "errors": 0, "chunks": 0}

    for i, pdf_info in enumerate(pdfs, start=1):
        pdf_path  = Path(pdf_info["full_path"])
        pdf_hash  = pdf_info["source_pdf_hash"]
        pdf_name  = pdf_info["filename"]

        if i % progress_interval == 0 or i == 1:
            logger.info(f"  再インデックス進捗: {i}/{total} — {pdf_name}")

        try:
            file_info = {
                "rel_path":        f"pdfs/{pdf_name}",
                "full_path":       str(pdf_path),
                "filename":        pdf_name,
                "category":        "reindex",
                "subcategory":     "",
                "sub_subcategory": "",
                "file_type":       "pdf",
                "file_size_kb":    float(pdf_path.stat().st_size) / 1024.0,
                "modified_at":     datetime.fromtimestamp(pdf_path.stat().st_mtime, tz=timezone.utc).isoformat(),
                "source_pdf_hash": pdf_hash,
                "source_pdf_name": pdf_name,
            }

            process_and_index_file(file_info, stats)

        except Exception as e:
            logger.error(f"  エラー ({pdf_name}): {e}", exc_info=True)
            stats["errors"] += 1

    logger.info(f"再インデックス完了: {stats}")
    return stats


# ─── 単一ファイルインデックス (外部API用) ────────────────────────────────────────
def index_file(filepath: str) -> Dict[str, Any]:
    """単一ファイルをインデックスに追加・更新"""
    path = Path(filepath)
    if not path.exists():
        return {"error": "File not found"}

    from config import KNOWLEDGE_BASE_DIR as _KB_DIR, EXCLUDE_FOLDERS, SUPPORTED_EXTENSIONS

    base_path = Path(_KB_DIR)
    try:
        rel_path = path.relative_to(base_path)
    except ValueError:
        return {"error": f"File is outside knowledge base dir: {path}"}

    parts = rel_path.parts
    category = parts[0] if len(parts) > 1 else ""

    file_info = {
        "filename":       path.name,
        "full_path":      str(path),
        "rel_path":       str(rel_path),
        "category":       category,
        "subcategory":    parts[1] if len(parts) > 2 else "",
        "sub_subcategory": parts[2] if len(parts) > 3 else "",
        "file_type":      path.suffix.lower().lstrip("."),
        "file_size_kb":   round(path.stat().st_size / 1024, 2),
        "modified_at":    datetime.fromtimestamp(path.stat().st_mtime, tz=timezone.utc).isoformat(),
        "doc_type":       _infer_doc_type(category, path.name),
    }

    client = get_chroma_client()
    embedding_function = GeminiEmbeddingFunction()
    collection = client.get_or_create_collection(
        name=COLLECTION_NAME,
        embedding_function=embedding_function,
    )

    try:
        collection.delete(where={"rel_path": str(rel_path)})
    except Exception as e:
        logger.warning(f"既存インデックス削除エラー (無視可能) ({rel_path}): {e}")

    stats = {"chunks": 0, "indexed": 0, "skipped": 0, "errors": 0}
    try:
        process_and_index_file(file_info, stats)
        stats["indexed"] = 1
        logger.info(f"単一ファイルインデックス完了: {stats['chunks']}チャンク ({path.name})")
    except Exception as e:
        logger.error(f"インデックスエラー ({path.name}): {e}", exc_info=True)
        return {"error": str(e)}

    return stats


def delete_from_index(rel_path: str) -> bool:
    """インデックスからファイルを削除"""
    try:
        client = get_chroma_client()
        collection = client.get_or_create_collection(name=COLLECTION_NAME)
        collection.delete(where={"rel_path": rel_path})
    except Exception as e:
        logger.error(f"ChromaDB削除エラー ({rel_path}): {e}", exc_info=True)

    try:
        result = _delete_doc_index(rel_path)
        logger.info(f"インデックス削除{'完了' if result else '（見つからず）'}: {rel_path}")
        return result
    except Exception as e:
        logger.error(f"DBインデックス削除エラー ({rel_path}): {e}", exc_info=True)
        return False


def delete_file_completely(file_path: str) -> dict:
    """
    ファイルに紐づく全データを完全削除する。
    - ChromaDB チャンク（rel_path & source_pdf_hash の両方で削除）
    - data/parent_chunks/ の JSON ファイル
    - data/pdfs/ のハッシュ名 PDF
    - files テーブル（file_store）
    - Document テーブル（SQLite）
    - knowledge_base 内の物理ファイル（PDF/MD）

    Args:
        file_path: knowledge_base からの相対パス（例: "06_設計/hoge.pdf"）

    Returns:
        dict: 削除結果サマリー
    """
    from pathlib import Path
    from config import KNOWLEDGE_BASE_DIR, PARENT_CHUNKS_DIR, PDF_STORAGE_DIR, COLLECTION_NAME
    import hashlib

    results: Dict[str, Any] = {
        "chroma_chunks": 0,
        "parent_chunks": 0,
        "pdf_storage": False,
        "file_store": False,
        "doc_index": False,
        "physical_files": [],
        "errors": [],
    }

    abs_path = Path(KNOWLEDGE_BASE_DIR) / file_path
    # PDF/MDどちらのパスが来ても両方を対象にする
    pdf_path = abs_path.with_suffix(".pdf") if abs_path.suffix.lower() == ".md" else abs_path
    md_path  = abs_path.with_suffix(".md")  if abs_path.suffix.lower() == ".pdf" else abs_path

    # --- 1. ChromaDB: rel_path で削除 ---
    try:
        client = get_chroma_client()
        col = client.get_or_create_collection(name=COLLECTION_NAME)

        # MDのrel_pathで削除（インデックスはMDで登録）
        md_rel = str(md_path.relative_to(Path(KNOWLEDGE_BASE_DIR)))
        existing = col.get(where={"rel_path": md_rel}, include=[])
        if existing["ids"]:
            col.delete(ids=existing["ids"])
            results["chroma_chunks"] += len(existing["ids"])
            logger.info(f"[Delete] ChromaDB rel_path削除: {len(existing['ids'])}件")
    except Exception as e:
        results["errors"].append(f"ChromaDB rel_path削除エラー: {e}")
        logger.warning(f"[Delete] ChromaDB rel_path削除エラー（続行）: {e}")

    # --- 2. ChromaDB: source_pdf_hash で削除（漏れチャンク対策） ---
    try:
        if pdf_path.exists():
            with open(pdf_path, "rb") as f:
                pdf_hash = hashlib.sha256(f.read()).hexdigest()
            existing = col.get(where={"source_pdf_hash": pdf_hash}, include=[])
            if existing["ids"]:
                col.delete(ids=existing["ids"])
                results["chroma_chunks"] += len(existing["ids"])
                logger.info(f"[Delete] ChromaDB hash削除: {len(existing['ids'])}件 (hash={pdf_hash[:8]}...)")
    except Exception as e:
        results["errors"].append(f"ChromaDB hash削除エラー: {e}")
        logger.warning(f"[Delete] ChromaDB hash削除エラー（続行）: {e}")

    # --- 3. parent_chunks JSON 削除 ---
    try:
        parent_dir = Path(PARENT_CHUNKS_DIR)
        if parent_dir.exists():
            # source_pdf_hash に基づくJSONを検索・削除
            if pdf_path.exists():
                deleted = 0
                for json_file in parent_dir.glob(f"{pdf_hash[:16]}*.json"):
                    json_file.unlink()
                    deleted += 1
                # ファイル名ベースでも検索（ハッシュが取れない場合のフォールバック）
                stem = pdf_path.stem
                for json_file in parent_dir.glob(f"*{stem}*.json"):
                    if json_file.exists():
                        json_file.unlink()
                        deleted += 1
                results["parent_chunks"] = deleted
                if deleted:
                    logger.info(f"[Delete] parent_chunks JSON削除: {deleted}件")
    except Exception as e:
        results["errors"].append(f"parent_chunks削除エラー: {e}")
        logger.warning(f"[Delete] parent_chunks削除エラー（続行）: {e}")

    # --- 4. data/pdfs/ のハッシュ名PDF削除 ---
    try:
        if pdf_path.exists():
            pdf_storage_dir = Path(PDF_STORAGE_DIR)
            hash_pdf = pdf_storage_dir / f"{pdf_hash}{pdf_path.suffix}"
            if hash_pdf.exists():
                hash_pdf.unlink()
                results["pdf_storage"] = True
                logger.info(f"[Delete] data/pdfs/ ハッシュPDF削除: {hash_pdf.name}")
    except Exception as e:
        results["errors"].append(f"data/pdfs/削除エラー: {e}")
        logger.warning(f"[Delete] data/pdfs/削除エラー（続行）: {e}")

    # --- 5. file_store (files テーブル) 削除 ---
    try:
        import file_store as fs
        # current_path で検索（PDF・MD両方試みる）
        for target_path in [str(pdf_path), str(md_path)]:
            rec = fs.get_file_by_path(target_path)
            if rec:
                fs.delete_file(rec["id"])
                results["file_store"] = True
                logger.info(f"[Delete] file_store削除: {rec['id']}")
                break
    except Exception as e:
        results["errors"].append(f"file_store削除エラー: {e}")
        logger.warning(f"[Delete] file_store削除エラー（続行）: {e}")

    # --- 6. Document テーブル（SQLite）削除 ---
    try:
        base_kb = Path(KNOWLEDGE_BASE_DIR)
        try:
            rel_md = str(md_path.relative_to(base_kb))
            _delete_doc_index(rel_md)
        except ValueError:
            pass
        
        _delete_doc_index(file_path)  # 元のパスでも試みる
        results["doc_index"] = True
    except Exception as e:
        results["errors"].append(f"Document削除エラー: {e}")
        logger.warning(f"[Delete] Document削除エラー（続行）: {e}")

    # --- 7. Status cleanup (Status management is now unified in MetadataRepository) ---
    try:
        # OCRStatusManager is deprecated and removed. 
        # Document deletion in step 6 already handles indexing status.
        pass
    except Exception as e:
        logger.warning(f"[Delete] Status cleanup error: {e}")

    # --- 8. 物理ファイル削除（knowledge_base内） ---
    for target in [pdf_path, md_path]:
        if target.exists():
            try:
                target.unlink()
                results["physical_files"].append(target.name)
                logger.info(f"[Delete] 物理ファイル削除: {target}")
            except Exception as e:
                results["errors"].append(f"物理ファイル削除エラー ({target.name}): {e}")
                logger.error(f"[Delete] 物理ファイル削除エラー: {target}: {e}")

    logger.info(
        f"[Delete] 完全削除完了: {file_path} | "
        f"chunks={results['chroma_chunks']}, "
        f"parent_chunks={results['parent_chunks']}, "
        f"files={results['physical_files']}"
    )
    return results


# ─── 後方互換 ────────────────────────────────────────────────────────────────────
def load_file_index() -> Dict[str, Any]:
    from database import get_session, LegacyDocument as DbDocument
    session = get_session()
    try:
        docs = session.query(DbDocument).filter(DbDocument.file_hash.isnot(None)).all()
        files = {}
        for doc in docs:
            files[doc.file_path] = {
                "hash":       doc.file_hash,
                "chunk_count": doc.chunk_count or 0,
                "indexed_at":  doc.last_indexed_at.isoformat() if doc.last_indexed_at else None,
                "modified_at": doc.updated_at.isoformat() if doc.updated_at else None,
            }
        last_updated = None
        if docs:
            valid = [d.last_indexed_at for d in docs if d.last_indexed_at]
            last_updated = max(valid).isoformat() if valid else None
        return {"files": files, "last_updated": last_updated}
    finally:
        session.close()


def save_file_index(index: Dict[str, Any]):
    pass  # DB で管理しているため不要


def chunk_text(text: str, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
    """後方互換シム: 旧 chunk_text を呼び出しているコード向け"""
    doc_type = metadata.get("doc_type", "catalog")
    source_pdf_hash = metadata.get("source_pdf", metadata.get("source_pdf_hash", "unknown"))
    source_pdf_name = metadata.get("source_pdf_name", metadata.get("filename", ""))
    return chunk_for_indexing(
        text=text,
        page_number=metadata.get("page_number"),
        has_image=False,
        doc_type=doc_type,
        source_pdf_hash=source_pdf_hash,
        source_pdf_name=source_pdf_name,
        category=metadata.get("category", ""),
        rel_path=metadata.get("rel_path", ""),
        filename=metadata.get("filename", ""),
        file_type=metadata.get("file_type", ""),
        modified_at=metadata.get("modified_at", ""),
        tags_str=metadata.get("tags_str", ""),
    )


if __name__ == "__main__":
    build_index()
